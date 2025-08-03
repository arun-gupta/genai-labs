import io
import json
from typing import Dict, Any, Optional
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown


class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def export_to_pdf(self, content: Dict[str, Any], filename: str = "generated_content.pdf") -> bytes:
        """Export content to PDF format."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        title = Paragraph("Generated Content", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Metadata
        if 'metadata' in content:
            metadata = content['metadata']
            metadata_data = [
                ['Field', 'Value'],
                ['Model Provider', metadata.get('model_provider', 'N/A')],
                ['Model Name', metadata.get('model_name', 'N/A')],
                ['Generated At', metadata.get('timestamp', 'N/A')],
                ['Token Usage', str(metadata.get('token_usage', 'N/A'))],
                ['Latency', f"{metadata.get('latency_ms', 0):.2f}ms" if metadata.get('latency_ms') else 'N/A'],
            ]
            
            metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(metadata_table)
            story.append(Spacer(1, 20))
        
        # System Prompt
        if 'system_prompt' in content and content['system_prompt']:
            story.append(Paragraph("System Prompt", self.styles['Heading2']))
            story.append(Paragraph(content['system_prompt'], self.styles['Normal']))
            story.append(Spacer(1, 12))
        
        # User Prompt
        if 'user_prompt' in content:
            story.append(Paragraph("User Prompt", self.styles['Heading2']))
            story.append(Paragraph(content['user_prompt'], self.styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Generated Content
        if 'generated_content' in content:
            story.append(Paragraph("Generated Content", self.styles['Heading2']))
            
            # Handle different content types
            generated_content = content['generated_content']
            if isinstance(generated_content, list):
                # Multiple candidates
                for i, candidate in enumerate(generated_content, 1):
                    story.append(Paragraph(f"Candidate {i}", self.styles['Heading3']))
                    story.append(Paragraph(str(candidate), self.styles['Normal']))
                    story.append(Spacer(1, 12))
            else:
                # Single content
                story.append(Paragraph(str(generated_content), self.styles['Normal']))
        
        # Analytics (if available)
        if 'analytics' in content and content['analytics']:
            story.append(Spacer(1, 20))
            story.append(Paragraph("Analytics Summary", self.styles['Heading2']))
            
            analytics = content['analytics']
            if 'generation_metrics' in analytics:
                metrics = analytics['generation_metrics']['basic_metrics']
                metrics_data = [
                    ['Metric', 'Value'],
                    ['Characters', str(metrics.get('characters', 'N/A'))],
                    ['Words', str(metrics.get('words', 'N/A'))],
                    ['Sentences', str(metrics.get('sentences', 'N/A'))],
                    ['Paragraphs', str(metrics.get('paragraphs', 'N/A'))],
                ]
                
                metrics_table = Table(metrics_data, colWidths=[2*inch, 2*inch])
                metrics_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(metrics_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_word(self, content: Dict[str, Any], filename: str = "generated_content.docx") -> bytes:
        """Export content to Word format."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Generated Content', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        if 'metadata' in content:
            doc.add_heading('Metadata', level=1)
            metadata = content['metadata']
            
            metadata_table = doc.add_table(rows=1, cols=2)
            metadata_table.style = 'Table Grid'
            hdr_cells = metadata_table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'
            
            metadata_items = [
                ('Model Provider', metadata.get('model_provider', 'N/A')),
                ('Model Name', metadata.get('model_name', 'N/A')),
                ('Generated At', metadata.get('timestamp', 'N/A')),
                ('Token Usage', str(metadata.get('token_usage', 'N/A'))),
                ('Latency', f"{metadata.get('latency_ms', 0):.2f}ms" if metadata.get('latency_ms') else 'N/A'),
            ]
            
            for field, value in metadata_items:
                row_cells = metadata_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = value
            
            doc.add_paragraph()
        
        # System Prompt
        if 'system_prompt' in content and content['system_prompt']:
            doc.add_heading('System Prompt', level=1)
            doc.add_paragraph(content['system_prompt'])
            doc.add_paragraph()
        
        # User Prompt
        if 'user_prompt' in content:
            doc.add_heading('User Prompt', level=1)
            doc.add_paragraph(content['user_prompt'])
            doc.add_paragraph()
        
        # Generated Content
        if 'generated_content' in content:
            doc.add_heading('Generated Content', level=1)
            
            generated_content = content['generated_content']
            if isinstance(generated_content, list):
                # Multiple candidates
                for i, candidate in enumerate(generated_content, 1):
                    doc.add_heading(f'Candidate {i}', level=2)
                    doc.add_paragraph(str(candidate))
                    doc.add_paragraph()
            else:
                # Single content
                doc.add_paragraph(str(generated_content))
        
        # Analytics (if available)
        if 'analytics' in content and content['analytics']:
            doc.add_heading('Analytics Summary', level=1)
            
            analytics = content['analytics']
            if 'generation_metrics' in analytics:
                metrics = analytics['generation_metrics']['basic_metrics']
                
                metrics_table = doc.add_table(rows=1, cols=2)
                metrics_table.style = 'Table Grid'
                hdr_cells = metrics_table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Value'
                
                metrics_items = [
                    ('Characters', str(metrics.get('characters', 'N/A'))),
                    ('Words', str(metrics.get('words', 'N/A'))),
                    ('Sentences', str(metrics.get('sentences', 'N/A'))),
                    ('Paragraphs', str(metrics.get('paragraphs', 'N/A'))),
                ]
                
                for metric, value in metrics_items:
                    row_cells = metrics_table.add_row().cells
                    row_cells[0].text = metric
                    row_cells[1].text = value
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_markdown(self, content: Dict[str, Any], filename: str = "generated_content.md") -> str:
        """Export content to Markdown format."""
        md_content = []
        
        # Title
        md_content.append("# Generated Content\n")
        
        # Metadata
        if 'metadata' in content:
            md_content.append("## Metadata\n")
            metadata = content['metadata']
            
            md_content.append("| Field | Value |")
            md_content.append("|-------|-------|")
            md_content.append(f"| Model Provider | {metadata.get('model_provider', 'N/A')} |")
            md_content.append(f"| Model Name | {metadata.get('model_name', 'N/A')} |")
            md_content.append(f"| Generated At | {metadata.get('timestamp', 'N/A')} |")
            md_content.append(f"| Token Usage | {metadata.get('token_usage', 'N/A')} |")
            latency = f"{metadata.get('latency_ms', 0):.2f}ms" if metadata.get('latency_ms') else 'N/A'
            md_content.append(f"| Latency | {latency} |")
            md_content.append("\n")
        
        # System Prompt
        if 'system_prompt' in content and content['system_prompt']:
            md_content.append("## System Prompt\n")
            md_content.append(content['system_prompt'])
            md_content.append("\n")
        
        # User Prompt
        if 'user_prompt' in content:
            md_content.append("## User Prompt\n")
            md_content.append(content['user_prompt'])
            md_content.append("\n")
        
        # Generated Content
        if 'generated_content' in content:
            md_content.append("## Generated Content\n")
            
            generated_content = content['generated_content']
            if isinstance(generated_content, list):
                # Multiple candidates
                for i, candidate in enumerate(generated_content, 1):
                    md_content.append(f"### Candidate {i}\n")
                    md_content.append(str(candidate))
                    md_content.append("\n")
            else:
                # Single content
                md_content.append(str(generated_content))
                md_content.append("\n")
        
        # Analytics (if available)
        if 'analytics' in content and content['analytics']:
            md_content.append("## Analytics Summary\n")
            
            analytics = content['analytics']
            if 'generation_metrics' in analytics:
                metrics = analytics['generation_metrics']['basic_metrics']
                
                md_content.append("| Metric | Value |")
                md_content.append("|--------|-------|")
                md_content.append(f"| Characters | {metrics.get('characters', 'N/A')} |")
                md_content.append(f"| Words | {metrics.get('words', 'N/A')} |")
                md_content.append(f"| Sentences | {metrics.get('sentences', 'N/A')} |")
                md_content.append(f"| Paragraphs | {metrics.get('paragraphs', 'N/A')} |")
                md_content.append("\n")
        
        return "\n".join(md_content)


# Global instance
export_service = ExportService() 